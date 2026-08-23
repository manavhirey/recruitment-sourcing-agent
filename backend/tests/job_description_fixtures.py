from io import BytesIO
from struct import pack
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject


def readable_docx() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_heading("Senior Product Designer", level=1)
    document.add_paragraph("Lead product design for the growth team.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Experience"
    table.cell(0, 1).text = "10+ years"
    document.save(output)
    return output.getvalue()


def interleaved_docx() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "First cell"
    table.cell(0, 1).text = "Second cell"
    document.add_paragraph("After table")
    document.save(output)
    return output.getvalue()


def empty_docx() -> bytes:
    output = BytesIO()
    Document().save(output)
    return output.getvalue()


def docx_with_text(text: str) -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(output)
    return output.getvalue()


def docx_with_external_relationship() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("Local job description text")
    document.part.relate_to(
        "https://invalid.example/job-description",
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    document.save(output)
    return output.getvalue()


def docx_with_malformed_document_xml() -> bytes:
    source = ZipFile(BytesIO(readable_docx()))
    output = BytesIO()
    with source, ZipFile(output, "w") as target:
        for entry in source.infolist():
            contents = source.read(entry)
            if entry.filename == "word/document.xml":
                contents = b"<w:document"
            target.writestr(entry, contents)
    return output.getvalue()


def docx_with_encrypted_member() -> bytes:
    data = bytearray(readable_docx())
    _patch_zip_member_headers(
        data,
        filename=b"word/document.xml",
        flag_bits=1,
    )
    return bytes(data)


def docx_with_unsupported_compression() -> bytes:
    data = bytearray(readable_docx())
    _patch_zip_member_headers(
        data,
        filename=b"word/document.xml",
        compression_method=99,
    )
    return bytes(data)


def docx_with_forged_small_expanded_size(actual_size: int) -> bytes:
    prefix = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:document xmlns:w="http://schemas.openxmlformats.org/'
        b'wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>'
    )
    suffix = b"</w:t></w:r></w:p></w:body></w:document>"
    padding_size = actual_size - len(prefix) - len(suffix)
    if padding_size < 0:
        raise ValueError("actual_size is too small for document XML")

    output = BytesIO()
    with ZipFile(BytesIO(readable_docx())) as source, ZipFile(output, "w") as target:
        for entry in source.infolist():
            if entry.filename != "word/document.xml":
                target.writestr(entry, source.read(entry))
                continue
            with target.open(entry, "w") as document_xml:
                document_xml.write(prefix)
                chunk = b"x" * 1_000_000
                remaining = padding_size
                while remaining:
                    portion = chunk[: min(remaining, len(chunk))]
                    document_xml.write(portion)
                    remaining -= len(portion)
                document_xml.write(suffix)

    data = bytearray(output.getvalue())
    _patch_central_file_size(data, b"word/document.xml", 1)
    return bytes(data)


def readable_pdf() -> bytes:
    return text_pdf("Senior Product Designer")


def text_pdf(text: str) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


def empty_pdf() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(output)
    return output.getvalue()


def encrypted_pdf() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("secret")
    writer.write(output)
    return output.getvalue()


def pdf_with_pages(count: int) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    for _ in range(count):
        writer.add_blank_page(width=612, height=792)
    writer.write(output)
    return output.getvalue()


def pdf_with_decoded_content_sizes(sizes: tuple[int, ...]) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    for size in sizes:
        page = writer.add_blank_page(width=612, height=792)
        stream = DecodedStreamObject()
        stream.set_data(b" " * size)
        page[NameObject("/Contents")] = writer._add_object(stream.flate_encode(9))
    writer.write(output)
    return output.getvalue()


def docx_package_with_entries(count: int) -> bytes:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as package:
        package.writestr("[Content_Types].xml", b"<Types />")
        package.writestr("word/document.xml", b"<document />")
        for index in range(count - 2):
            package.writestr(f"padding/{index}.txt", b"")
    return output.getvalue()


def docx_with_declared_expanded_bytes(size: int) -> bytes:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as package:
        package.writestr("[Content_Types].xml", b"<Types />")
        package.writestr("word/document.xml", b"<document />")
        package.writestr("padding.bin", b"x")

    data = bytearray(output.getvalue())
    _patch_central_file_size(data, b"padding.bin", size)
    return bytes(data)


def _patch_central_file_size(data: bytearray, filename: bytes, size: int) -> None:
    marker = b"PK\x01\x02"
    offset = 0
    while (offset := data.find(marker, offset)) != -1:
        name_length = int.from_bytes(data[offset + 28 : offset + 30], "little")
        current_filename = bytes(data[offset + 46 : offset + 46 + name_length])
        if current_filename == filename:
            data[offset + 24 : offset + 28] = pack("<I", size)
            return
        offset += 4
    raise AssertionError(f"{filename!r} central-directory record not found")


def _patch_zip_member_headers(
    data: bytearray,
    *,
    filename: bytes,
    flag_bits: int | None = None,
    compression_method: int | None = None,
) -> None:
    records = (
        (b"PK\x03\x04", 26, 30, 6, 8),
        (b"PK\x01\x02", 28, 46, 8, 10),
    )
    for (
        marker,
        name_length_offset,
        name_offset,
        flag_offset,
        compression_offset,
    ) in records:
        offset = 0
        while (offset := data.find(marker, offset)) != -1:
            name_length = int.from_bytes(
                data[offset + name_length_offset : offset + name_length_offset + 2],
                "little",
            )
            current_filename = bytes(
                data[offset + name_offset : offset + name_offset + name_length]
            )
            if current_filename == filename:
                if flag_bits is not None:
                    data[offset + flag_offset : offset + flag_offset + 2] = pack(
                        "<H", flag_bits
                    )
                if compression_method is not None:
                    data[
                        offset + compression_offset : offset + compression_offset + 2
                    ] = pack("<H", compression_method)
                break
            offset += 4
        else:
            raise AssertionError(f"{filename!r} ZIP record not found")


def docx_with_nested_archive(
    *,
    extension: str = ".zip",
    signature: bytes = b"PK\x03\x04",
) -> bytes:
    output = BytesIO(readable_docx())
    output.seek(0, 2)
    with ZipFile(output, "a", ZIP_DEFLATED) as package:
        package.writestr(f"word/embeddings/nested{extension}", signature + b"nested")
    return output.getvalue()
